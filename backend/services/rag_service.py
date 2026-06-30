import re
import shutil
import unicodedata
from pathlib import Path

import chromadb
import pdfplumber
from docx import Document as DocxDocument
from llama_index.core import Settings, StorageContext, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import Document
from llama_index.vector_stores.chroma import ChromaVectorStore
from pdf2image import convert_from_path
from rapidocr_onnxruntime import RapidOCR

from backend.config import (
    CHILD_CHUNK_OVERLAP,
    CHILD_CHUNK_SIZE,
    CHROMA_DIR,
    COLLECTION_NAME,
    EMBED_MODEL,
    LLM_MODEL,
    MODEL_PROVIDER,
    PARENT_CHUNK_OVERLAP,
    PARENT_CHUNK_SIZE,
    SIMILARITY_TOP_K,
    SOURCE_SCORE_THRESHOLD,
)


NO_ANSWER = "Không tìm thấy thông tin này trong tài liệu."
LLM_FALLBACK_ANSWER = (
    "Không gọi được model để tạo câu trả lời hoàn chỉnh. "
    "Dưới đây là các đoạn tài liệu liên quan nhất."
)
MAX_VISIBLE_SOURCES = 4
SOURCE_PREVIEW_LENGTH = 180
SOURCE_STOPWORDS = {
    "cua", "cac", "cho", "voi", "trong", "ngoai", "nhung", "duoc", "the",
    "nay", "kia", "la", "va", "hoac", "thi", "mot", "ve", "tu", "den",
    "what", "when", "where", "which", "that", "this", "are", "the",
}


def setup_llamaindex():
    provider = MODEL_PROVIDER.lower().strip()

    if provider == "ollama":
        from llama_index.embeddings.ollama import OllamaEmbedding
        from llama_index.llms.ollama import Ollama

        Settings.embed_model = OllamaEmbedding(model_name=EMBED_MODEL)
        Settings.llm = Ollama(model=LLM_MODEL, request_timeout=180.0)
    elif provider == "openai":
        from llama_index.embeddings.openai import OpenAIEmbedding
        from llama_index.llms.openai import OpenAI

        Settings.embed_model = OpenAIEmbedding(model_name=EMBED_MODEL)
        Settings.llm = OpenAI(model=LLM_MODEL, request_timeout=180.0)
    else:
        # Fallback to OpenAI if the selected provider is not available.
        from llama_index.embeddings.openai import OpenAIEmbedding
        from llama_index.llms.openai import OpenAI

        Settings.embed_model = OpenAIEmbedding(model_name=EMBED_MODEL)
        Settings.llm = OpenAI(model=LLM_MODEL, request_timeout=180.0)

    Settings.text_splitter = SentenceSplitter(
        chunk_size=CHILD_CHUNK_SIZE,
        chunk_overlap=CHILD_CHUNK_OVERLAP,
    )


def normalize_for_match(text: str) -> str:
    text = unicodedata.normalize("NFD", (text or "").lower())
    text = "".join(char for char in text if unicodedata.category(char) != "Mn")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def is_no_answer(answer: str) -> bool:
    normalized = normalize_for_match(answer)
    return (
        "khong tim thay" in normalized
        or "toi khong tim thay" in normalized
        or "khong co thong tin" in normalized
        or "khong du thong tin" in normalized
    )


def is_good_text(text: str) -> bool:
    if not text or len(text.strip()) < 50:
        return False

    readable_chars = sum(
        1 for char in text
        if char.isalnum() or char.isspace() or char in ".,;:!?()[]{}+-=*/_@#$%&<>"
    )

    return readable_chars / max(len(text), 1) > 0.75


def normalize_source_text(text: str) -> str:
    text = re.sub(r"---\s*Trang\s+\d+(?:\s+OCR)?\s*---", " ", text or "")
    return re.sub(r"\s+", " ", text).strip()


def read_pdf_pages_with_pdfplumber(file_path: Path):
    page_items = []

    with pdfplumber.open(str(file_path)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = (page.extract_text() or "").strip()

            if page_text:
                page_items.append({
                    "text": page_text,
                    "page": page_number,
                    "reader": "pdfplumber",
                })

    return page_items


def read_pdf_pages_with_ocr(file_path: Path):
    ocr = RapidOCR()
    pages = convert_from_path(str(file_path), dpi=200)
    page_items = []

    for page_number, image in enumerate(pages, start=1):
        result, _ = ocr(image)
        page_lines = []

        if result:
            for line in result:
                page_lines.append(line[1])

        page_text = "\n".join(page_lines).strip()

        if page_text:
            page_items.append({
                "text": page_text,
                "page": page_number,
                "reader": "ocr",
            })

    return page_items


def read_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore").strip()


def read_docx(file_path: Path) -> str:
    doc = DocxDocument(str(file_path))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return "\n".join(paragraphs)


def read_document_pages(file_path: Path):
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        pages = read_pdf_pages_with_pdfplumber(file_path)
        combined_text = "\n\n".join(page["text"] for page in pages)

        if is_good_text(combined_text):
            return pages

        print("PDF extract failed, switching to OCR...")
        return read_pdf_pages_with_ocr(file_path)

    if suffix == ".docx":
        return [{"text": read_docx(file_path), "page": None, "reader": "docx"}]

    if suffix == ".txt":
        return [{"text": read_txt(file_path), "page": None, "reader": "txt"}]

    raise ValueError("Chỉ hỗ trợ PDF, DOCX và TXT")


def chunk_text(text: str, chunk_size: int, overlap: int):
    clean_text = normalize_source_text(text)

    if not clean_text:
        return []

    if len(clean_text) <= chunk_size:
        return [{
            "text": clean_text,
            "start": 0,
            "end": len(clean_text),
        }]

    chunks = []
    start = 0
    step = max(chunk_size - overlap, 1)

    while start < len(clean_text):
        end = min(start + chunk_size, len(clean_text))
        chunks.append({
            "text": clean_text[start:end].strip(),
            "start": start,
            "end": end,
        })

        if end >= len(clean_text):
            break

        start += step

    return chunks


def parent_for_child(child, parent_chunks):
    child_center = (child["start"] + child["end"]) // 2

    for parent in parent_chunks:
        if parent["start"] <= child_center <= parent["end"]:
            return parent

    return parent_chunks[0] if parent_chunks else child


def build_parent_child_documents(file_path: Path, page_items):
    documents = []

    for section_index, item in enumerate(page_items, start=1):
        section_text = item["text"].strip()

        if not section_text:
            continue

        parent_chunks = chunk_text(
            section_text,
            PARENT_CHUNK_SIZE,
            PARENT_CHUNK_OVERLAP,
        )
        child_chunks = chunk_text(
            section_text,
            CHILD_CHUNK_SIZE,
            CHILD_CHUNK_OVERLAP,
        )

        for child_index, child in enumerate(child_chunks, start=1):
            parent = parent_for_child(child, parent_chunks)
            metadata = {
                "file_name": file_path.name,
                "file_path": str(file_path),
                "reader": item["reader"],
                "section": section_index,
                "chunk_id": f"{section_index}-{child_index}",
                "child_text": child["text"],
                "parent_context": parent["text"],
            }

            if item["page"] is not None:
                metadata["page"] = item["page"]

            documents.append(Document(
                text=child["text"],
                metadata=metadata,
                excluded_embed_metadata_keys=[
                    "file_path",
                    "reader",
                    "section",
                    "chunk_id",
                    "child_text",
                    "parent_context",
                ],
                excluded_llm_metadata_keys=[
                    "file_path",
                    "reader",
                    "section",
                    "chunk_id",
                    "child_text",
                    "parent_context",
                ],
            ))

    return documents


def question_keywords(question: str):
    words = re.findall(r"[\wÀ-ỹ]+", (question or "").lower(), flags=re.UNICODE)
    keywords = []

    for word in words:
        normalized = normalize_for_match(word)

        if len(normalized) < 3 or normalized in SOURCE_STOPWORDS or normalized in keywords:
            continue

        keywords.append(normalized)

    return keywords


def make_source_preview(text: str, question: str) -> str:
    clean_text = normalize_source_text(text)

    if len(clean_text) <= SOURCE_PREVIEW_LENGTH:
        return clean_text

    normalized_text = normalize_for_match(clean_text)
    hit_index = -1

    for keyword in question_keywords(question):
        hit_index = normalized_text.find(keyword)
        if hit_index >= 0:
            break

    if hit_index < 0:
        hit_index = 0

    start = max(hit_index - 60, 0)
    end = min(start + SOURCE_PREVIEW_LENGTH, len(clean_text))
    snippet = clean_text[start:end].strip()

    if start > 0:
        snippet = f"...{snippet}"
    if end < len(clean_text):
        snippet = f"{snippet}..."

    return snippet


def score_is_usable(score):
    return score is None or score >= SOURCE_SCORE_THRESHOLD


def build_sources(nodes, question: str):
    sources = []
    seen_locations = set()

    for node in nodes:
        metadata = getattr(node, "metadata", None) or {}
        score = float(node.score) if node.score is not None else None

        if not score_is_usable(score):
            continue

        file_name = metadata.get("file_name") or "Tài liệu"
        page_number = metadata.get("page")
        chunk_id = metadata.get("chunk_id")
        location_key = (file_name, page_number, chunk_id)
        context = metadata.get("parent_context") or node.text or ""
        child_text = metadata.get("child_text") or node.text or context

        if location_key in seen_locations:
            continue

        preview = make_source_preview(child_text, question)

        if not preview:
            continue

        seen_locations.add(location_key)
        sources.append({
            "index": len(sources) + 1,
            "file_name": file_name,
            "page": page_number,
            "section": metadata.get("section"),
            "chunk_id": chunk_id,
            "score": score,
            "preview": preview,
            "context": normalize_source_text(context),
        })

        if len(sources) >= MAX_VISIBLE_SOURCES:
            break

    return sources


def build_context_from_sources(sources):
    context_blocks = []

    for source in sources:
        location = source["file_name"]

        if source["page"]:
            location = f"{location}, trang {source['page']}"
        elif source["section"]:
            location = f"{location}, đoạn {source['section']}"

        context_blocks.append(
            f"[{source['index']}] {location}\n{source['context']}"
        )

    return "\n\n".join(context_blocks)


def public_sources(sources):
    return [
        {
            "index": source["index"],
            "file_name": source["file_name"],
            "page": source["page"],
            "section": source["section"],
            "chunk_id": source["chunk_id"],
            "score": source["score"],
            "preview": source["preview"],
        }
        for source in sources
    ]


def get_chroma_collection():
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)

    client = chromadb.PersistentClient(path=str(CHROMA_DIR))

    return client.get_or_create_collection(COLLECTION_NAME)


def get_index():
    setup_llamaindex()

    collection = get_chroma_collection()
    vector_store = ChromaVectorStore(chroma_collection=collection)

    return VectorStoreIndex.from_vector_store(
        vector_store=vector_store,
        embed_model=Settings.embed_model,
    )


def index_document(file_path: Path):
    setup_llamaindex()

    if not file_path.exists():
        raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

    page_items = read_document_pages(file_path)
    text = "\n\n".join(item["text"] for item in page_items if item["text"].strip())

    print("=" * 60)
    print("FILE:", file_path.name)
    print("SECTIONS:", len(page_items))
    print("TEXT LENGTH:", len(text))
    print("TEXT PREVIEW:")
    print(text[:3000])
    print("=" * 60)

    if not text or len(text.strip()) < 20:
        raise ValueError("Không đọc được nội dung tài liệu. File có thể là PDF scan ảnh cần OCR.")

    documents = build_parent_child_documents(file_path, page_items)

    if not documents:
        raise ValueError("Không tạo được chunk nào từ tài liệu.")

    collection = get_chroma_collection()
    vector_store = ChromaVectorStore(chroma_collection=collection)

    storage_context = StorageContext.from_defaults(
        vector_store=vector_store
    )

    VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        embed_model=Settings.embed_model,
        show_progress=True,
    )

    return {
        "document_count": len(documents),
        "section_count": len(page_items),
        "text_length": len(text),
        "vector_database": "ChromaDB",
        "rag_framework": "LlamaIndex",
        "chunking": "parent-child",
        "parent_chunk_size": PARENT_CHUNK_SIZE,
        "child_chunk_size": CHILD_CHUNK_SIZE,
        "reader": page_items[0]["reader"] if page_items else None,
        "embedding_model": EMBED_MODEL,
        "llm_model": LLM_MODEL,
    }


def ask_document(question: str):
    setup_llamaindex()

    index = get_index()
    retriever = index.as_retriever(
        similarity_top_k=SIMILARITY_TOP_K
    )
    nodes = retriever.retrieve(question)

    if not nodes:
        return {
            "answer": NO_ANSWER,
            "sources": [],
        }

    sources = build_sources(nodes, question)

    if not sources:
        return {
            "answer": NO_ANSWER,
            "sources": [],
        }

    context = build_context_from_sources(sources)
    prompt = f"""
Bạn là chatbot hỏi đáp tài liệu.

Chỉ sử dụng thông tin trong phần Tài liệu để trả lời.

Quy tắc:
- Trả lời trực tiếp câu hỏi.
- Trả lời tự nhiên bằng tiếng Việt.
- Khi dùng thông tin từ nguồn nào, thêm mã nguồn dạng [1], [2] ngay sau ý tương ứng.
- Chỉ dùng các mã nguồn có trong phần Tài liệu.
- Nếu tài liệu không có thông tin đủ rõ để trả lời, chỉ trả lời:
  "{NO_ANSWER}"

Tài liệu:
{context}

Câu hỏi:
{question}
"""

    try:
        response = Settings.llm.complete(prompt)
        answer = str(response).strip()
    except Exception as error:
        print(f"LLM generation failed: {error}")
        return {
            "answer": LLM_FALLBACK_ANSWER,
            "sources": public_sources(sources),
        }

    if is_no_answer(answer):
        return {
            "answer": NO_ANSWER,
            "sources": [],
        }

    return {
        "answer": answer,
        "sources": public_sources(sources),
    }


def clear_vector_database():
    if CHROMA_DIR.exists():
        shutil.rmtree(CHROMA_DIR)

    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
